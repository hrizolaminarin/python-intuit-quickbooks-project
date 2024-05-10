# Install python-docx, a package that allows us to work with documents, using pip3, package manager for python used to install existing packages that other developers developed
# by typing the following command in VS Code's Terminal:
# pip3 install python-docx
from docx import Document
from docx.shared import Inches
# Install pyttsx3, python Text to Speach library for converting text to speach:
# pip3 install pyttsx3
# 
# Both packages from above could also be installed using requirements.txt file with the following command:
# pip3 install -r requirements.txt
# Requirements.txt file is usefull for defining all the dependencies (specify package name and version in a following format: python-docx==1.1.2)
#
# To uninstall package we can use the following command:
# pip3 uninstall python-docx
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# Profile picture taken from pexels.com
document.add_picture('me.jpg', width=Inches(2.0))

# Name Phone and Email details
speak('What is your name? ')
name = input('What is your name? ')
speak('Hello ' + name + ' how are you today?')
speak('What is your phone number? ')
phone_number = input('What is your phone number? ')
speak('What is your email? ')
email = input('What is your email? ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email
)

# About me
document.add_heading('About me')
speak('Tell me about yourself? ')
about_me = input('Tell me about yourself? ')
document.add_paragraph(about_me)

# Work experience
document.add_heading('Work experience')
p = document.add_paragraph()

speak('Enter company ')
company = input('Enter company ')
speak('From Date ')
from_date = input('From Date ')
speak('To Date ')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True

speak('Describe your experience at ' + company + ' ')
experience_details =  input('Describe your experience at ' + company + ' ')
p.add_run(experience_details)

# More experiences
while True:
    speak('Do you have more experiences? Yes or No ')
    has_more_experience = input('Do you have more experiences? Yes or No ')
    if has_more_experience.lower() == 'yes':
        p = document.add_paragraph()

        speak('Enter company ')
        company = input('Enter company ')
        speak('From Date ')
        from_date = input('From Date ')
        speak('To Date ')
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True

        speak('Describe your experience at ' + company + ' ')
        experience_details =  input('Describe your experience at ' + company + ' ')
        p.add_run(experience_details)
    else:
        break

# Skills
document.add_heading('Skills')

speak('Let us know your skills ')
skill = input('Let us know your skills ')
document.add_paragraph(skill, 'List Bullet')

while True:
    speak('Do you have more skills? Yes or No ')
    has_more_skills = input('Do you have more skills? Yes or No ')
    if has_more_skills.lower() == 'yes':
        speak('Let us know your skill ')
        skill = input('Let us know your skill ')
        document.add_paragraph(skill, 'List Bullet')
    else:
        break


# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Amigoscode and Intuit QuickBooks course project"

document.save('cv.docx')